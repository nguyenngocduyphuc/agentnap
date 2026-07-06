#!/usr/bin/env python3
"""Generate AgentNap promo one-pagers: VN + EN, .docx + .pdf.

Run with the workspace venv (needs python-docx, reportlab):
  .venv/bin/python make_collateral.py
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

OUT = Path(__file__).parent
BRAND = "1A4D7A"  # deep blue
REPO = "github.com/nguyenngocduyphuc/agentnap"

VN = {
    "file": "AgentNap_GioiThieu_VN",
    "title": "AgentNap 😴 — Đòi lại RAM mà AI coding agent đang rò rỉ trên máy bạn",
    "sections": [
        ("Vấn đề: máy 32GB vẫn nghẹt thở", [
            "Claude Code, Codex, Cursor và các MCP server sinh ra hàng trăm process con. "
            "Khi phiên làm việc đóng không sạch, chúng thành process 'mồ côi' — mỗi cái giữ "
            "200MB–1GB RAM vĩnh viễn cho tới khi khởi động lại máy.",
            "Đây không phải lỗi của riêng bạn: kho issue của claude-code ghi nhận máy bị chiếm "
            "12GB, 60GB, thậm chí 108GB RAM (issues #23252, #18859, #56960). Máy tác giả từng có "
            "~400 process agent, swap 26.5GB, quạt gào cả ngày.",
        ]),
        ("Giải pháp: dọn đúng rác, không đụng việc đang chạy", [
            "AgentNap là công cụ dòng lệnh MIỄN PHÍ, mã nguồn mở, KHÔNG cần cài thêm thư viện:",
            "• agentnap advise — chẩn đoán bằng ngôn ngữ thường: cái gì đang ăn RAM, cái gì dọn "
            "được ngay, cái gì cần bạn quyết.",
            "• agentnap reap — dọn CHỈ process mồ côi (cha đã chết), tắt êm trước rồi mới ép. "
            "Mặc định chạy thử (dry-run), không giết gì khi chưa có --apply.",
            "• agentnap daemon — trực nền, chỉ hành động khi macOS báo áp lực bộ nhớ thật.",
            "• agentnap stats — 'biên lai' tổng RAM đã đòi lại được.",
            "• advise --ai — kèm kế hoạch cá nhân hóa từ API AI bạn tự chọn (DeepSeek/OpenAI/"
            "Ollama...), key của bạn, dữ liệu không đi đâu khác.",
            "CAM KẾT: hành động tự động duy nhất là dọn process có cha đã chết, đang idle, "
            "không phải app GUI. Mọi thứ có thể đụng phiên làm việc sống chỉ là LỜI KHUYÊN.",
        ]),
        ("Đã kiểm định độc lập — không chỉ là lời nói", [
            "• Audit an toàn bởi model AI độc lập (DeepSeek): 4/4 tuyên bố an toàn ĐẠT; "
            "các phát hiện đã được sửa trước khi phát hành.",
            "• Thí nghiệm tái lập được: tạo 5 orphan thật + 1 process đang hoạt động — bắt đúng "
            "5/5, process đang hoạt động còn nguyên. Thí nghiệm này chạy tự động trên máy chủ "
            "sạch của GitHub (macOS + Windows) ở MỖI lần cập nhật code.",
            "• Kết quả thực tế trên máy tác giả trong 1 ngày: swap 26.5GB → 5.2GB, memory "
            "pressure WARNING → normal, 0 phiên làm việc bị gián đoạn.",
        ]),
        ("Dùng thử trong 30 giây", [
            "git clone https://" + REPO + " && cd agentnap && ./install.sh",
            "agentnap advise   →   agentnap reap --apply   →   agentnap stats",
            "macOS đầy đủ; Windows bản beta (advise/status/reap). Bản Pro (app thanh menu, "
            "tự động hoàn toàn) đang mở danh sách chờ tại Issue #1 trên GitHub.",
        ]),
        ("Tác giả", [
            "Phạm Đức Phương Nam — kỹ sư tự động hóa dược phẩm (GxP), vận hành song song 9+ AI "
            "agent mỗi ngày; AgentNap sinh ra từ chính nhu cầu đó. "
            "Web: phamducphuongnam.com · nampham.net · Repo: " + REPO,
        ]),
    ],
}

EN = {
    "file": "AgentNap_Overview_EN",
    "title": "AgentNap 😴 — Reclaim the RAM your AI coding agents are leaking",
    "sections": [
        ("The problem: 32GB machines gasping for air", [
            "Claude Code, Codex, Cursor and MCP servers spawn hundreds of child processes. "
            "When sessions close uncleanly they become orphans — each holding 200MB–1GB of RAM "
            "until you reboot.",
            "It is not just you: the claude-code issue tracker documents machines at 12GB, 60GB, "
            "even 108GB (issues #23252, #18859, #56960). The author's Mac hit ~400 agent "
            "processes and 26.5GB of swap.",
        ]),
        ("The fix: clean the garbage, never touch live work", [
            "AgentNap is a FREE, open-source, zero-dependency CLI:",
            "• agentnap advise — plain-language diagnosis: what eats RAM, what is safe to "
            "reclaim now, what is your call.",
            "• agentnap reap — removes ONLY orphaned processes (parent already dead), gracefully "
            "(SIGTERM → grace → SIGKILL). Dry-run by default.",
            "• agentnap daemon — background watchdog that acts only on real macOS "
            "memory-pressure signals.",
            "• agentnap stats — receipts: total RAM reclaimed so far.",
            "• advise --ai — optional personalized plan from YOUR OWN LLM key "
            "(DeepSeek/OpenAI/Ollama — any OpenAI-compatible API).",
            "THE GUARANTEE: the only automatic action is reaping orphaned, idle, non-GUI agent "
            "processes. Anything that could touch a live session is surfaced as advice.",
        ]),
        ("Independently verified — not just claimed", [
            "• Safety audit by an independent AI model (DeepSeek): 4/4 safety claims PASS; "
            "findings fixed before release.",
            "• Repeatable experiment: 5 real orphans + 1 active look-alike — 5/5 detected and "
            "reaped, the active process untouched. Runs in CI on clean GitHub macOS AND Windows "
            "runners on every push.",
            "• Field result on the author's machine in one day: swap 26.5GB → 5.2GB, memory "
            "pressure WARNING → normal, zero interrupted sessions.",
        ]),
        ("Try it in 30 seconds", [
            "git clone https://" + REPO + " && cd agentnap && ./install.sh",
            "agentnap advise   →   agentnap reap --apply   →   agentnap stats",
            "Full support on macOS; Windows beta (advise/status/reap). Pro (menu-bar app, fully "
            "automatic) waitlist: GitHub Issue #1.",
        ]),
        ("Author", [
            "Nam Pham (Phạm Đức Phương Nam) — pharma GxP automation engineer running 9+ AI "
            "agents in parallel daily; AgentNap was built to survive that. "
            "Web: phamducphuongnam.com · nampham.net · Repo: " + REPO,
        ]),
    ],
}


def make_docx(doc_spec: dict) -> Path:
    d = Document()
    style = d.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)
    h = d.add_heading(doc_spec["title"], level=0)
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(BRAND)
        run.font.size = Pt(17)
    for head, paras in doc_spec["sections"]:
        hh = d.add_heading(head, level=1)
        for run in hh.runs:
            run.font.color.rgb = RGBColor.from_string(BRAND)
            run.font.size = Pt(13)
        for p in paras:
            d.add_paragraph(p)
    out = OUT / f"{doc_spec['file']}.docx"
    d.save(out)
    return out


def make_pdf(doc_spec: dict) -> Path:
    pdfmetrics.registerFont(
        TTFont("ArialUnicode", "/Library/Fonts/Arial Unicode.ttf"))
    title_s = ParagraphStyle("t", fontName="ArialUnicode", fontSize=16,
                             leading=21, textColor=HexColor("#" + BRAND),
                             spaceAfter=6)
    head_s = ParagraphStyle("h", fontName="ArialUnicode", fontSize=12.5,
                            leading=16, textColor=HexColor("#" + BRAND),
                            spaceBefore=10, spaceAfter=3)
    body_s = ParagraphStyle("b", fontName="ArialUnicode", fontSize=10,
                            leading=14.5, spaceAfter=4)
    out = OUT / f"{doc_spec['file']}.pdf"
    doc = SimpleDocTemplate(str(out), pagesize=A4,
                            leftMargin=18 * mm, rightMargin=18 * mm,
                            topMargin=16 * mm, bottomMargin=16 * mm,
                            title=doc_spec["title"], author="Nam Pham")
    story = [Paragraph(doc_spec["title"], title_s), Spacer(1, 4)]
    for head, paras in doc_spec["sections"]:
        story.append(Paragraph(head, head_s))
        story.extend(Paragraph(p, body_s) for p in paras)
    doc.build(story)
    return out


if __name__ == "__main__":
    for spec in (VN, EN):
        print(make_docx(spec))
        print(make_pdf(spec))
